# Youtube video Analysis 
# example url: https://www.youtube.com/watch?v=X8MZWCGgIb8

import gradio as gr
import yt_dlp as youtube_dl
from moviepy.editor import AudioFileClip
from pydub import AudioSegment
import os
import nltk
import re
from rake_nltk import Rake
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.datasets import fetch_20newsgroups
from sklearn.model_selection import train_test_split
from sklearn.pipeline import Pipeline
from sklearn.naive_bayes import MultinomialNB
import numpy as np
from docx import Document
import speech_recognition as sr
from docx.shared import Pt

def create_report(extracted_text, processed_text, keywords_rake, keywords_tfidf, predicted_category, recommendations):
    doc = Document()
    doc.add_heading('YouTube Video Analysis Report', 0)
    doc.add_heading('Extracted Text', level=1)
    doc.add_paragraph(extracted_text)
    doc.add_heading('Processed Text', level=1)
    doc.add_paragraph(processed_text)
    doc.add_heading('Keywords (RAKE)', level=1)
    if keywords_rake:
        doc.add_paragraph(', '.join(keywords_rake))
    else:
        doc.add_paragraph('No keywords extracted.')
    doc.add_heading('Keywords (TF-IDF)', level=1)
    if keywords_tfidf:
        for keyword, score in keywords_tfidf:
            doc.add_paragraph(f"{keyword}: {score:.4f}")
    else:
        doc.add_paragraph('No keywords extracted.')
    doc.add_heading('Predicted Category', level=1)
    doc.add_paragraph(predicted_category)
    doc.add_heading('Recommendations', level=1)
    if recommendations:
        for rec, score in recommendations:
            doc.add_paragraph(f"Similarity: {score:.4f}\n{rec[:500]}...") 
    else:
        doc.add_paragraph('No recommendations available.')
    doc.save('YouTube_Video_Analysis_Report.docx')

# Download and process the YouTube audio
def process_youtube_audio(url):
    ydl_opts = {
        'format': 'bestaudio/best',
        'outtmpl': 'audio.%(ext)s',
        'postprocessors': [{
            'key': 'FFmpegExtractAudio',
            'preferredcodec': 'wav',
            'preferredquality': '192',
        }],
    }
    with youtube_dl.YoutubeDL(ydl_opts) as ydl:
        ydl.download([url])

    audio_path = 'audio_output.wav'
    audio_clip = AudioFileClip("audio.wav")
    audio_clip.write_audiofile(audio_path)

    def split_audio(file_path, chunk_length_ms=60000):
        audio = AudioSegment.from_wav(file_path)
        chunks = [audio[i:i + chunk_length_ms] for i in range(0, len(audio), chunk_length_ms)]
        chunk_files = []
        
        for i, chunk in enumerate(chunks):
            chunk_filename = f"chunk_{i}.wav"
            chunk.export(chunk_filename, format="wav")
            chunk_files.append(chunk_filename)
        
        return chunk_files

    # Split the audio into chunks
    chunks = split_audio("audio_output.wav")

    # Function to recognize speech with retry
    def recognize_speech_with_retry(audio_file, retries=3, delay=5):
        recognizer = sr.Recognizer()
        for attempt in range(retries):
            try:
                with sr.AudioFile(audio_file) as source:
                    audio = recognizer.record(source)
                    text = recognizer.recognize_google(audio)
                    return text
            except sr.RequestError as e:
                if attempt < retries - 1:
                    time.sleep(delay)
                else:
                    return "Recognition error"
            except Exception as e:
                return "Error occurred"

    # Process each chunk individually
    all_text = []
    for chunk in chunks:
        text = recognize_speech_with_retry(chunk)
        if text:
            all_text.append(text)
        os.remove(chunk)  # Clean up the chunk file after processing

    final_text = " ".join(all_text)
    return final_text

# Text Preprocessing
def preprocess_text(text):
    text = text.lower()  # Lowercase text
    text = re.sub(r'\d+', '', text)  # Remove numbers
    text = re.sub(r'\W', ' ', text)  # Remove punctuation
    text = re.sub(r'\s+', ' ', text)  # Remove extra spaces
    tokens = nltk.word_tokenize(text)  # Tokenization
    stopwords = set(nltk.corpus.stopwords.words('english'))
    tokens = [word for word in tokens if word not in stopwords]  # Remove stopwords
    return " ".join(tokens)

# Keyword Extraction
def extract_keywords(text):
    r = Rake()
    r.extract_keywords_from_text(text)
    keywords = r.get_ranked_phrases()
    return keywords

def extract_keywords_tfidf(text):
    vectorizer = TfidfVectorizer(stop_words='english')
    X = vectorizer.fit_transform([text])
    feature_names = vectorizer.get_feature_names_out()
    scores = X.T.sum(axis=1).A1
    keywords = dict(zip(feature_names, scores))
    sorted_keywords = sorted(keywords.items(), key=lambda x: x[1], reverse=True)
    return sorted_keywords

# Classification
def classify_text(text):
    categories = ['alt.atheism', 'soc.religion.christian', 'comp.graphics', 'sci.med']
    newsgroups = fetch_20newsgroups(subset='train', categories=categories)
    X_train, X_test, y_train, y_test = train_test_split(newsgroups.data, newsgroups.target, test_size=0.2, random_state=42)
    model = Pipeline([('tfidf', TfidfVectorizer()), ('clf', MultinomialNB())])
    model.fit(X_train, y_train)
    predicted_category = model.predict([text])
    return newsgroups.target_names[predicted_category[0]]

# Recommendation System
def recommend_videos(text):
    categories = ['alt.atheism', 'soc.religion.christian', 'comp.graphics', 'sci.med']
    newsgroups = fetch_20newsgroups(subset='train', categories=categories)
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform([text] + newsgroups.data)
    cos_sim = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix)
    similarities = cos_sim[0][1:]  # Exclude the first comparison with itself
    top_indices = np.argsort(similarities)[-5:][::-1]  # Get top 5 most similar
    recommended_videos = [(newsgroups.data[i], similarities[i]) for i in top_indices]
    return [(video[:100] + "...", sim) for video, sim in recommended_videos]

def gradio_interface(url):
    # Process YouTube audio and extract text
    text = process_youtube_audio(url)
    
    # Preprocess the extracted text
    processed_text = preprocess_text(text)
    
    # Extract keywords using RAKE
    keywords_rake = extract_keywords(processed_text)
    
    # Extract keywords using TF-IDF
    keywords_tfidf = extract_keywords_tfidf(processed_text)
    
    # Classify the text
    category = classify_text(processed_text)
    
    # Recommend similar content
    recommendations = recommend_videos(processed_text)
    
    create_report(text, processed_text, keywords_rake, keywords_tfidf, category, recommendations)

    # Return outputs as a tuple
    return text, processed_text, keywords_rake, keywords_tfidf, category, recommendations

# Define the Gradio interface
iface = gr.Interface(
    fn=gradio_interface,
    inputs="text",
    outputs=[
        gr.Textbox(label="Extracted Text"),
        gr.Textbox(label="Processed Text"),
        gr.JSON(label="Keywords (RAKE)"),
        gr.JSON(label="Keywords (TF-IDF)"),
        gr.Textbox(label="Predicted Category"),
        gr.JSON(label="Recommendations")
    ],
    live=True,
    title="YouTube Video Analysis",
    description="Process a YouTube video URL to extract text, keywords, classify content, and recommend similar videos."
)

iface.launch()